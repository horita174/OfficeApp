from docx import Document
from docx.shared import Inches

# ドキュメントを読み込む
doc = Document("sample.docx")

# 画像を追加する
doc.add_picture("./img/doge.png", width=Inches(4), height=Inches(3))

# テキストの文字数をカウントする
total_characters = 0
for paragraph in doc.paragraphs:
    total_characters += len(paragraph.text)

# 結果を表示する
doc.add_paragraph("文字数: " + str(total_characters))

# 変更を保存する
doc.save("sample_answer.docx")